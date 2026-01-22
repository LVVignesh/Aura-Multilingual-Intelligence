import gradio as gr
import speech_recognition as sr
from deep_translator import GoogleTranslator
from gtts import gTTS
import tempfile
import os
from pydub import AudioSegment
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.corpus import stopwords
from nltk.probability import FreqDist
from heapq import nlargest
import nltk
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document as DocxDocument
import io

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)
    
try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords', quiet=True)

# Initialize translator
# translator = Translator()  # Not needed for deep-translator as we instantiate per call or use a helper

# Language mapping
LANGUAGES = {
    'English': 'en', 'Spanish': 'es', 'French': 'fr', 'German': 'de',
    'Italian': 'it', 'Portuguese': 'pt', 'Russian': 'ru', 'Japanese': 'ja',
    'Korean': 'ko', 'Chinese (Simplified)': 'zh-cn', 'Hindi': 'hi',
    'Arabic': 'ar', 'Dutch': 'nl', 'Greek': 'el', 'Turkish': 'tr',
    'Vietnamese': 'vi', 'Thai': 'th', 'Polish': 'pl', 'Danish': 'da',
    'Finnish': 'fi', 'Czech': 'cs', 'Romanian': 'ro', 'Hungarian': 'hu',
    'Swedish': 'sv', 'Indonesian': 'id', 'Hebrew': 'he', 'Bengali': 'bn',
    'Kannada': 'kn'
}

def transcribe_audio(audio_file):
    """Transcribe audio file to text"""
    if audio_file is None:
        return "Please upload an audio file or record audio."
    
    recognizer = sr.Recognizer()
    
    try:
        # Convert to WAV if needed
        audio = AudioSegment.from_file(audio_file)
        with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as tmp_wav:
            wav_path = tmp_wav.name
            audio.export(wav_path, format='wav')
        
        # Transcribe
        with sr.AudioFile(wav_path) as source:
            recognizer.adjust_for_ambient_noise(source, duration=0.5)
            audio_data = recognizer.record(source)
            text = recognizer.recognize_google(audio_data)
        
        # Cleanup
        os.remove(wav_path)
        return text
        
    except sr.UnknownValueError:
        return "Could not understand audio"
    except sr.RequestError as e:
        return f"API Error: {str(e)}"
    except Exception as e:
        return f"Error: {str(e)}"

def translate_text(text, target_language):
    """Translate text to target language"""
    if not text or not text.strip():
        return "Please provide text to translate."
    
    try:
        target_lang = LANGUAGES.get(target_language, 'en')
        if target_lang == 'en' and target_language == 'English':
            return text  # No translation needed
        
        # Check if source is auto
        translated = GoogleTranslator(source='auto', target=target_lang).translate(text)
        return translated
    except Exception as e:
        return f"Translation error: {str(e)}"

def text_to_speech(text, language):
    """Convert text to speech"""
    if not text or not text.strip():
        return None
    
    try:
        lang_code = LANGUAGES.get(language, 'en')
        tts = gTTS(text=text, lang=lang_code, slow=False)
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.mp3') as tmp_audio:
            audio_path = tmp_audio.name
            tts.save(audio_path)
        
        return audio_path
    except Exception as e:
        print(f"TTS Error: {e}")
        return None

def extract_keywords(text, num_keywords=5):
    """Extract keywords from text"""
    if not text or not text.strip():
        return "Please provide text to analyze."
    
    try:
        stop_words = set(stopwords.words('english'))
        words = word_tokenize(text.lower())
        words = [word for word in words if word.isalnum() and word not in stop_words]
        
        if not words:
            return "No keywords found."
        
        freq_dist = FreqDist(words)
        keywords = nlargest(num_keywords, freq_dist, key=freq_dist.get)
        
        return "\n".join([f"{i+1}. {word}" for i, word in enumerate(keywords)])
    except Exception as e:
        return f"Error: {str(e)}"

def generate_summary(text, num_sentences=3):
    """Generate summary from text"""
    if not text or not text.strip():
        return "Please provide text to summarize."
    
    try:
        stop_words = set(stopwords.words('english'))
        sentences = sent_tokenize(text)
        
        # Handle short text
        if len(sentences) <= num_sentences:
            return text
        
        # If very long unpunctuated text, split by word count
        if len(sentences) <= 2 and len(text.split()) > 50:
            words_list = text.split()
            chunk_size = max(15, len(words_list) // (num_sentences * 2))
            sentences = []
            for i in range(0, len(words_list), chunk_size):
                chunk = ' '.join(words_list[i:i+chunk_size])
                if chunk.strip():
                    sentences.append(chunk)
        
        # Calculate word frequencies
        words = word_tokenize(text.lower())
        words = [word for word in words if word.isalnum() and word not in stop_words]
        
        if not words:
            return text[:200] + "..." if len(text) > 200 else text
        
        word_freq = FreqDist(words)
        sentence_scores = {}
        
        # Score sentences
        for sentence in sentences:
            words_in_sentence = word_tokenize(sentence.lower())
            word_count = len([word for word in words_in_sentence if word.isalnum()])
            
            if word_count == 0:
                continue
            
            score = sum(word_freq[word.lower()] for word in words_in_sentence if word.lower() in word_freq)
            score = score / max(5, word_count)
            sentence_scores[sentence] = score
        
        if not sentence_scores:
            return text[:200] + "..." if len(text) > 200 else text
        
        # Select top sentences
        num_to_select = min(num_sentences, len(sentences))
        summary_sentences = nlargest(num_to_select, sentence_scores.items(), key=lambda x: x[1])
        summary_sentences.sort(key=lambda x: text.find(x[0]))
        
        return " ".join([s[0] for s in summary_sentences])
    except Exception as e:
        return f"Error: {str(e)}"

def save_as_txt(text):
    """Save text as TXT file"""
    if not text or not text.strip():
        return None
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.txt', mode='w', encoding='utf-8') as tmp:
        tmp.write(text)
        return tmp.name

def save_as_markdown(text):
    """Save text as Markdown file"""
    if not text or not text.strip():
        return None
    
    markdown_content = f"""# Automatic Notes Maker Output
**Generated on:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

---

{text}
"""
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.md', mode='w', encoding='utf-8') as tmp:
        tmp.write(markdown_content)
        return tmp.name

def save_as_pdf(text):
    """Save text as PDF file"""
    if not text or not text.strip():
        return None
    
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
            pdf_path = tmp.name
        
        doc = SimpleDocTemplate(pdf_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#2196F3',
            spaceAfter=30
        )
        story.append(Paragraph("Automatic Notes Maker", title_style))
        story.append(Spacer(1, 12))
        
        # Date
        date_text = f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        story.append(Paragraph(date_text, styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Content
        for paragraph in text.split('\n'):
            if paragraph.strip():
                story.append(Paragraph(paragraph, styles['Normal']))
                story.append(Spacer(1, 12))
        
        doc.build(story)
        return pdf_path
    except Exception as e:
        print(f"PDF Error: {e}")
        return None

def save_as_docx(text):
    """Save text as Word document"""
    if not text or not text.strip():
        return None
    
    try:
        doc = DocxDocument()
        doc.add_heading('Automatic Notes Maker', 0)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph('')
        
        for paragraph in text.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph)
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            return tmp.name
    except Exception as e:
        print(f"DOCX Error: {e}")
        return None

# Gradio Interface Functions
def process_audio_and_translate(audio, target_language, auto_translate):
    """Process audio: transcribe and optionally translate"""
    transcription = transcribe_audio(audio)
    
    if auto_translate and target_language != "English":
        translation = translate_text(transcription, target_language)
    else:
        translation = ""
    
    return transcription, translation

def translate_and_speak(text, target_language):
    """Translate text and generate speech"""
    translation = translate_text(text, target_language)
    audio = text_to_speech(translation, target_language)
    return translation, audio

def analyze_text(text, num_keywords, num_sentences):
    """Analyze text: extract keywords and generate summary"""
    keywords = extract_keywords(text, num_keywords)
    summary = generate_summary(text, num_sentences)
    return keywords, summary

# Create Gradio Interface
with gr.Blocks(theme=gr.themes.Soft(), title="Automatic Notes Maker") as demo:
    gr.Markdown("""
    # 🎙️ Automatic Notes Maker
    ### Record, Transcribe, Translate, and Analyze Audio in 28+ Languages
    """)
    
    with gr.Tabs():
        # Tab 1: Live Recording & Translation
        with gr.Tab("🎙️ Live Recording"):
            gr.Markdown("### Record or upload audio, transcribe, and translate")
            
            with gr.Row():
                with gr.Column():
                    audio_input = gr.Audio(sources=["microphone", "upload"], type="filepath", label="Record or Upload Audio")
                    target_lang = gr.Dropdown(choices=list(LANGUAGES.keys()), value="Hindi", label="Translate to")
                    auto_translate_check = gr.Checkbox(value=True, label="Auto-translate")
                    process_btn = gr.Button("🎯 Process Audio", variant="primary")
                
                with gr.Column():
                    transcription_output = gr.Textbox(label="Original Text (Transcription)", lines=8)
                    translation_output = gr.Textbox(label="Translated Text", lines=8)
            
            with gr.Row():
                speak_btn = gr.Button("🔊 Speak Translation")
                audio_output = gr.Audio(label="Text-to-Speech Output", type="filepath")
            
            process_btn.click(
                fn=process_audio_and_translate,
                inputs=[audio_input, target_lang, auto_translate_check],
                outputs=[transcription_output, translation_output]
            )
            
            speak_btn.click(
                fn=lambda text, lang: text_to_speech(text, lang),
                inputs=[translation_output, target_lang],
                outputs=audio_output
            )
        
        # Tab 2: Text Translation
        with gr.Tab("🌍 Translation"):
            gr.Markdown("### Translate text to any language")
            
            with gr.Row():
                with gr.Column():
                    input_text = gr.Textbox(label="Input Text", lines=10, placeholder="Enter text to translate...")
                    trans_target_lang = gr.Dropdown(choices=list(LANGUAGES.keys()), value="Hindi", label="Translate to")
                    translate_btn = gr.Button("🔄 Translate", variant="primary")
                
                with gr.Column():
                    translated_text = gr.Textbox(label="Translated Text", lines=10)
                    tts_audio = gr.Audio(label="Listen to Translation", type="filepath")
            
            translate_btn.click(
                fn=translate_and_speak,
                inputs=[input_text, trans_target_lang],
                outputs=[translated_text, tts_audio]
            )
        
        # Tab 3: Analysis
        with gr.Tab("🧠 Analysis"):
            gr.Markdown("### Extract keywords and generate summaries")
            
            with gr.Row():
                with gr.Column():
                    analysis_input = gr.Textbox(label="Input Text", lines=10, placeholder="Paste or type text to analyze...")
                    with gr.Row():
                        num_keywords = gr.Slider(minimum=3, maximum=15, value=5, step=1, label="Number of Keywords")
                        num_summary_sentences = gr.Slider(minimum=1, maximum=10, value=3, step=1, label="Summary Sentences")
                    analyze_btn = gr.Button("📊 Analyze", variant="primary")
                
                with gr.Column():
                    keywords_output = gr.Textbox(label="Keywords", lines=8)
                    summary_output = gr.Textbox(label="Summary", lines=8)
            
            analyze_btn.click(
                fn=analyze_text,
                inputs=[analysis_input, num_keywords, num_summary_sentences],
                outputs=[keywords_output, summary_output]
            )
        
        # Tab 4: Export
        with gr.Tab("💾 Export"):
            gr.Markdown("### Save your work in various formats")
            
            export_text = gr.Textbox(label="Text to Export", lines=10, placeholder="Paste text to export...")
            
            with gr.Row():
                txt_btn = gr.Button("📄 Save as TXT")
                md_btn = gr.Button("📝 Save as Markdown")
                pdf_btn = gr.Button("📕 Save as PDF")
                docx_btn = gr.Button("📘 Save as Word")
            
            export_output = gr.File(label="Download File")
            
            txt_btn.click(fn=save_as_txt, inputs=export_text, outputs=export_output)
            md_btn.click(fn=save_as_markdown, inputs=export_text, outputs=export_output)
            pdf_btn.click(fn=save_as_pdf, inputs=export_text, outputs=export_output)
            docx_btn.click(fn=save_as_docx, inputs=export_text, outputs=export_output)
    
    gr.Markdown("""
    ---
    ### 🌟 Features
    - **28+ Languages**: Translate to Hindi, Spanish, French, German, Japanese, and more
    - **Smart Analysis**: Extract keywords and generate summaries using NLP
    - **Multi-Format Export**: Save as TXT, PDF, Word, or Markdown
    - **Text-to-Speech**: Listen to translations in any language
    
    ### 💡 Tips
    - For best results, speak clearly when recording
    - Auto-translate works in real-time for non-English languages
    - Use the Analysis tab to extract key insights from long transcripts
    """)

# Launch the app
if __name__ == "__main__":
    demo.launch()
